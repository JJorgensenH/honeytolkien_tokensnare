from docx import Document
from docx.shared import Inches
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml import CT_Inline, CT_Picture
from datetime import datetime

from .common import register_token, random_creation_date, random_modification_date

def inject_tracking_pixel(paragraph, tracking_url):
    """
    Inyecta una imagen externa invisible (1x1 px) manipulando el XML del documento.
    """
    run = paragraph.add_run()
    
    # 1. Crear la relación externa (el link al servidor)
    rels = run.part.rels
    rel_id = rels._next_rId
    rels.add_relationship(
        reltype=RELATIONSHIP_TYPE.IMAGE,
        target=tracking_url,
        rId=rel_id,
        is_external=True
    )

    # 2. Construir el XML de la imagen (CT_Picture)
    # Definimos un tamaño diminuto (0.01 pulgadas)
    cx = Inches(0.01)
    cy = Inches(0.01) 

    # XML hack para insertar la imagen vinculada
    pic_xml = CT_Picture._pic_xml() # Plantilla base
    pic = parse_xml(pic_xml)
    
    # Configurar IDs y Nombre
    pic.nvPicPr.cNvPr.id = 0
    pic.nvPicPr.cNvPr.name = 'tracking_pixel.png'
    
    # Vincular con la relación externa creada arriba
    pic.blipFill.blip.link = rel_id

    # Configurar tamaño en el XML
    pic.spPr.cx = cx
    pic.spPr.cy = cy

    # 3. Insertar en el párrafo
    shape_id = run.part.next_id
    inline = CT_Inline.new(cx, cy, shape_id, pic)
    run._r.add_drawing(inline)


def generate_docx_honeytoken(server_url, output_file, description, title=None, author=None, content=None):
    """
    Genera un archivo .docx con un pixel de tracking externo y metadatos anti-forense.
    """
    token_data = register_token(
        server_url,
        token_type="docx",
        description=description
    )
    
    tracking_url = token_data['tracking_url_image']
    
    # common.py devuelve strings ISO, pero python-docx necesita objetos datetime
    str_created = random_creation_date()
    str_modified = random_modification_date(str_created)
    dt_created = datetime.strptime(str_created, '%Y-%m-%dT%H:%M:%SZ')
    dt_modified = datetime.strptime(str_modified, '%Y-%m-%dT%H:%M:%SZ')

    # Crea el Documento
    doc = Document()

    # Contenido visible
    if title:
        doc.add_heading(title, level=1)
    if content:
        doc.add_paragraph(content)

    # Inyecta el pixel de tracking
    p = doc.add_paragraph()
    inject_tracking_pixel(p, tracking_url)

    # Agregamos Metadatos
    core = doc.core_properties
    core.title = title if title else ""
    core.author = author if author else ""
    core.created = dt_created
    core.modified = dt_modified

    doc.save(output_file)